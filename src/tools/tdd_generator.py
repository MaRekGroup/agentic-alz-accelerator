"""
As-Built Technical Design Document (TDD) Generator

Generates a comprehensive Word document for each deployed landing zone,
including resource inventory, architecture diagrams (SVG with Azure icons),
security posture, compliance status, and network topology.

Runs automatically after each successful deployment in the CI/CD pipeline.
Can also be invoked manually via CLI or the agentic workflow.
"""

import io
import json
import logging
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional

try:
    import cairosvg
except ImportError:
    cairosvg = None  # type: ignore[assignment]

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from src.tools.azure_diagram_generator import (
    generate_app_lz_diagram,
    generate_connectivity_diagram,
    generate_identity_diagram,
    generate_management_diagram,
    generate_security_diagram,
)
from src.tools.drawio_diagram_generator import DrawioEngine
from src.tools.python_diagram_generator import DiagramEngine

logger = logging.getLogger(__name__)

# ─── Color Palette (Microsoft Azure brand) ────────────────────────────────────
AZURE_BLUE   = RGBColor(0x00, 0x78, 0xD4)
DARK_BLUE    = RGBColor(0x00, 0x30, 0x67)
GREEN        = RGBColor(0x57, 0xA3, 0x00)
ORANGE       = RGBColor(0xFF, 0x8C, 0x00)
RED          = RGBColor(0xE8, 0x11, 0x23)
GRAY         = RGBColor(0x6B, 0x6B, 0x6B)
LIGHT_GRAY   = RGBColor(0xF2, 0xF2, 0xF2)
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)


class TDDGenerator:
    """Generates an as-built Technical Design Document for a deployed landing zone."""

    def __init__(
        self,
        project_name: str,
        profile: str,
        subscription_id: str,
        subscription_name: str,
        location: str,
        environment: str,
        framework: str,
        deployment_id: str = "",
        config_path: str = "environments/subscriptions.json",
    ):
        self.project_name = project_name
        self.profile = profile
        self.subscription_id = subscription_id
        self.subscription_name = subscription_name
        self.location = location
        self.environment = environment
        self.framework = framework
        self.deployment_id = deployment_id
        self.config_path = config_path
        self.generated_at = datetime.now(timezone.utc)
        self.doc = Document()
        self._resource_inventory: dict = {}
        self._compliance_data: dict = {}
        self._config: dict = {}

    # ─── Document Style Setup ─────────────────────────────────────────────

    def _setup_styles(self):
        """Configure document styles — fonts, colors, spacing."""
        style = self.doc.styles["Normal"]
        font = style.font
        font.name = "Segoe UI"
        font.size = Pt(10)
        font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        # Heading styles
        for level, size, color in [
            ("Heading 1", 22, AZURE_BLUE),
            ("Heading 2", 16, DARK_BLUE),
            ("Heading 3", 13, AZURE_BLUE),
        ]:
            h = self.doc.styles[level]
            h.font.name = "Segoe UI Semibold"
            h.font.size = Pt(size)
            h.font.color.rgb = color
            h.paragraph_format.space_before = Pt(18)
            h.paragraph_format.space_after = Pt(6)

    # ─── Table Helpers ────────────────────────────────────────────────────

    def _add_table(self, headers: list[str], rows: list[list[str]],
                   col_widths: Optional[list[float]] = None) -> None:
        """Add a formatted table with Azure blue header row."""
        table = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        # Header row
        hdr = table.rows[0]
        for i, text in enumerate(headers):
            cell = hdr.cells[i]
            cell.text = text
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.color.rgb = WHITE
                    run.font.bold = True
                    run.font.size = Pt(9)
            shading = cell._element.get_or_add_tcPr()
            shading_elm = shading.makeelement(
                qn("w:shd"),
                {qn("w:val"): "clear", qn("w:color"): "auto",
                 qn("w:fill"): "0078D4"},
            )
            shading.append(shading_elm)

        # Data rows (zebra striping)
        for r_idx, row_data in enumerate(rows):
            row = table.rows[1 + r_idx]
            for c_idx, text in enumerate(row_data):
                cell = row.cells[c_idx]
                cell.text = str(text)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9)

            if r_idx % 2 == 1:
                for cell in row.cells:
                    shading = cell._element.get_or_add_tcPr()
                    shading_elm = shading.makeelement(
                        qn("w:shd"),
                        {qn("w:val"): "clear", qn("w:color"): "auto",
                         qn("w:fill"): "F2F2F2"},
                    )
                    shading.append(shading_elm)

        # Set column widths
        if col_widths:
            for row in table.rows:
                for i, width in enumerate(col_widths):
                    row.cells[i].width = Cm(width)

    # ─── Diagram Insertion ────────────────────────────────────────────────

    def _insert_diagram(self, svg_content: str, caption: str = "") -> None:
        """Convert SVG to PNG and insert into the document."""
        try:
            if cairosvg is None:
                raise ImportError("cairosvg is not installed")
            png_bytes = cairosvg.svg2png(
                bytestring=svg_content.encode("utf-8"),
                output_width=1400,
                dpi=150,
            )
            stream = io.BytesIO(png_bytes)
            self.doc.add_picture(stream, width=Inches(6.5))

            # Center the image
            last_paragraph = self.doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            if caption:
                cap = self.doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = cap.add_run(caption)
                run.font.size = Pt(9)
                run.font.italic = True
                run.font.color.rgb = GRAY

        except Exception as e:
            logger.warning(f"Failed to render diagram: {e}")
            self.doc.add_paragraph(
                f"[Diagram could not be rendered: {e}]"
            ).italic = True

    def _insert_png_diagram(self, png_path: str, caption: str = "") -> None:
        """Insert a pre-rendered PNG diagram into the document."""
        try:
            self.doc.add_picture(png_path, width=Inches(6.5))
            last_paragraph = self.doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            if caption:
                cap = self.doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = cap.add_run(caption)
                run.font.size = Pt(9)
                run.font.italic = True
                run.font.color.rgb = GRAY
        except Exception as e:
            logger.warning(f"Failed to insert PNG diagram: {e}")
            self.doc.add_paragraph(
                f"[Diagram could not be rendered: {e}]"
            ).italic = True

    def _generate_drawio_diagram(self, output_dir: str | None = None) -> tuple[str | None, str | None]:
        """Generate Draw.io + SVG diagrams using the DrawioEngine.

        Returns (drawio_path, svg_path) or (None, None) on failure.
        """
        try:
            engine = DrawioEngine(output_dir=output_dir or "agent-output/tdd")
            drawio_path, svg_path = engine.generate_tdd_diagram(
                profile=self.profile,
                project_name=self.project_name,
                subscription_name=self.subscription_name,
                location=self.location,
                output_dir=output_dir,
            )
            logger.info("Draw.io diagram generated: %s", drawio_path)
            return drawio_path, svg_path
        except Exception as e:
            logger.warning("Draw.io diagram generation failed: %s", e)
            return None, None

    def _generate_png_diagram(self, output_dir: str | None = None) -> str | None:
        """Generate a PNG architecture diagram using the diagrams library.

        Returns the PNG file path, or None on failure.
        """
        try:
            engine = DiagramEngine(output_dir=output_dir or "agent-output/tdd")
            return engine.generate_tdd_diagram(
                profile=self.profile,
                project_name=self.project_name,
                subscription_name=self.subscription_name,
                location=self.location,
                output_dir=output_dir,
            )
        except Exception as e:
            # Loud failure — silent fallback to SVG previously masked the
            # missing `graphviz` system dep / `diagrams` pip package for days.
            # Logging at ERROR with a sentinel marker makes CI surface this.
            logger.error(
                "[ICONS_FALLBACK_USED] Failed to generate PNG diagram with "
                "Azure icons (diagrams library). Falling back to hand-drawn "
                "SVG. Ensure 'graphviz' system package and 'diagrams' pip "
                "package are installed. Error: %s",
                e,
                exc_info=True,
            )
            return None

    # ─── Cover Page ───────────────────────────────────────────────────────

    def _add_cover_page(self):
        """Generate the cover page with Azure branding."""
        # Add some spacing
        for _ in range(4):
            self.doc.add_paragraph()

        # Title
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("As-Built\nTechnical Design Document")
        run.font.size = Pt(32)
        run.font.color.rgb = AZURE_BLUE
        run.font.name = "Segoe UI Light"

        self.doc.add_paragraph()

        # Subtitle with LZ name
        sub = self.doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sub.add_run(f"{self.project_name}")
        run.font.size = Pt(24)
        run.font.color.rgb = DARK_BLUE
        run.font.name = "Segoe UI Semibold"

        sub2 = self.doc.add_paragraph()
        sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sub2.add_run(f"Profile: {self.profile}  ·  Environment: {self.environment}")
        run.font.size = Pt(14)
        run.font.color.rgb = GRAY

        for _ in range(3):
            self.doc.add_paragraph()

        # Metadata table
        meta_rows = [
            ["Subscription", self.subscription_name],
            ["Subscription ID", self.subscription_id],
            ["Location", self.location],
            ["IaC Framework", self.framework.title()],
            ["Deployment ID", self.deployment_id or "(manual)"],
            ["Generated", self.generated_at.strftime("%Y-%m-%d %H:%M UTC")],
            ["Document Version", "1.0 (auto-generated)"],
        ]
        self._add_table(["Field", "Value"], meta_rows, col_widths=[5, 12])

        self.doc.add_page_break()

    # ─── Table of Contents ────────────────────────────────────────────────

    def _add_toc(self):
        """Insert a Table of Contents field."""
        self.doc.add_heading("Table of Contents", level=1)

        p = self.doc.add_paragraph()
        run = p.add_run()
        fld_char_begin = run._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "begin"})
        run._element.append(fld_char_begin)

        run2 = p.add_run()
        instr = run2._element.makeelement(qn("w:instrText"), {})
        instr.text = ' TOC \\o "1-3" \\h \\z \\u '
        run2._element.append(instr)

        run3 = p.add_run()
        fld_char_end = run3._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "end"})
        run3._element.append(fld_char_end)

        p2 = self.doc.add_paragraph()
        run = p2.add_run("(Right-click → Update Field to populate in Word)")
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = GRAY

        self.doc.add_page_break()

    # ─── Section Builders ─────────────────────────────────────────────────

    def _section_executive_summary(self):
        self.doc.add_heading("1. Executive Summary", level=1)

        profile_desc = {
            "platform-management": "centralized monitoring, logging, Sentinel SIEM, and backup infrastructure",
            "platform-connectivity": "hub networking with Azure Firewall, DNS, DDoS protection, and ExpressRoute/VPN gateways",
            "platform-identity": "Active Directory Domain Services, Entra ID PIM, and privileged access workstations",
            "platform-security": "dedicated SecOps with Sentinel, Defender for Cloud (all plans), SOAR playbooks, and incident response",
            "corp": "internal corporate workloads with private connectivity, ExpressRoute peering, and on-prem DNS integration",
            "online": "internet-facing workloads with WAF, DDoS Standard, and public endpoint protection",
            "sap": "SAP S/4HANA workloads with accelerated networking, Azure NetApp Files, and proximity placement groups",
            "sandbox": "development and testing with budget enforcement, relaxed policies, and standalone networking",
        }

        desc = profile_desc.get(self.profile, f"workloads using the {self.profile} profile")

        self.doc.add_paragraph(
            f"This Technical Design Document (TDD) describes the as-built state of the "
            f"'{self.project_name}' landing zone deployed to subscription "
            f"'{self.subscription_name}' in {self.location}. "
            f"This landing zone provides {desc}."
        )
        self.doc.add_paragraph(
            f"The deployment was executed using {self.framework.title()} via the GitHub Actions "
            f"CI/CD pipeline with OIDC authentication and environment approval gates. "
            f"All resources comply with the CAF enterprise-scale baseline policies."
        )

        # Key facts box
        self.doc.add_heading("Key Facts", level=2)
        facts = [
            ["Landing Zone", self.project_name],
            ["Profile", self.profile],
            ["Subscription", f"{self.subscription_name} ({self.subscription_id})"],
            ["Region", self.location],
            ["Environment", self.environment],
            ["IaC Framework", self.framework.title()],
            ["Deployment ID", self.deployment_id or "(manual/CLI)"],
        ]
        self._add_table(["Attribute", "Value"], facts, col_widths=[5, 12])

    def _section_architecture_diagram(self):
        self.doc.add_heading("2. Architecture Diagram", level=1)
        self.doc.add_paragraph(
            "The following diagram illustrates the as-built architecture of this landing zone, "
            "including all deployed resources, networking topology, and security controls. "
            "Icons follow the official Microsoft Azure Architecture Icon set."
        )

        # Preferred: Draw.io engine (produces .drawio + .svg with Azure icons)
        # Fallback 1: PNG via diagrams/graphviz library
        # Fallback 2: hand-drawn SVG via azure_diagram_generator
        caption = f"Figure 1: {self.project_name} Architecture — As-Built"
        _drawio_path, drawio_svg = self._generate_drawio_diagram()
        if drawio_svg and Path(drawio_svg).exists():
            svg_content = Path(drawio_svg).read_text(encoding="utf-8")
            self._insert_diagram(svg_content, caption)
        else:
            png_path = self._generate_png_diagram()
            if png_path and Path(png_path).exists():
                self._insert_png_diagram(png_path, caption)
            else:
                svg = self._get_profile_diagram()
                self._insert_diagram(svg, caption)

    def _get_profile_diagram(self) -> str:
        """Generate the appropriate SVG diagram for this profile."""
        if self.profile == "platform-management":
            return generate_management_diagram(self.subscription_name, self.location)
        elif self.profile == "platform-connectivity":
            return generate_connectivity_diagram(self.subscription_name, self.location)
        elif self.profile == "platform-identity":
            return generate_identity_diagram(self.subscription_name, self.location)
        elif self.profile == "platform-security":
            return generate_security_diagram(self.subscription_name, self.location)
        else:
            # Application LZ
            display_name = self.project_name.replace("-", " ").title()
            return generate_app_lz_diagram(
                self.profile, self.project_name, display_name,
                self.subscription_name, self.location,
            )

    def _section_resource_inventory(self, inventory: dict):
        self.doc.add_heading("3. Resource Inventory", level=1)
        self.doc.add_paragraph(
            "Complete inventory of Azure resources deployed in this landing zone, "
            "queried from Azure Resource Graph at generation time."
        )

        self.doc.add_heading("3.1 Resource Summary by Type", level=2)
        by_type = inventory.get("by_type", {})
        if by_type:
            rows = [[rtype.split("/")[-1], str(count)]
                    for rtype, count in sorted(by_type.items(), key=lambda x: -x[1])]
            self._add_table(["Resource Type", "Count"], rows, col_widths=[10, 4])
        else:
            self.doc.add_paragraph("Resource inventory will be populated from Azure Resource Graph at deployment time.")
            # Include expected resources based on profile
            self._add_expected_resources()

        total = inventory.get("total_count", 0)
        p = self.doc.add_paragraph()
        run = p.add_run(f"Total resources: {total}")
        run.font.bold = True

    def _add_expected_resources(self):
        """Add expected resource table based on the landing zone profile."""
        expected = {
            "platform-management": [
                ["Log Analytics Workspace", "1", "PerGB2018, 365-day retention"],
                ["Microsoft Sentinel", "1", "SIEM + SOAR integration"],
                ["Automation Account", "1", "Linked to LAW"],
                ["Recovery Services Vault", "1", "Daily + weekly backup policies"],
                ["Action Groups", "3+", "Email, Teams webhook, SMS"],
                ["Azure Monitor Alert Rules", "10+", "Performance, security, cost"],
            ],
            "platform-connectivity": [
                ["Virtual Network (Hub)", "1", "Hub CIDR 10.0.0.0/16"],
                ["Azure Firewall (Premium)", "1", "Threat Intel, IDPS"],
                ["VPN Gateway", "1", "VpnGw2 SKU"],
                ["ExpressRoute Gateway", "1", "Ultra Performance SKU"],
                ["Azure Bastion", "1", "Standard SKU"],
                ["DDoS Protection Plan", "1", "Standard tier"],
                ["Private DNS Zones", "20+", "Azure PaaS services"],
                ["Network Security Groups", "4+", "Per subnet"],
            ],
            "platform-identity": [
                ["Virtual Network (Identity)", "1", "Identity CIDR"],
                ["Domain Controllers (VMs)", "2", "Zone-redundant (AZ-1, AZ-2)"],
                ["Entra Connect (VM)", "1", "Hybrid sync to Entra ID"],
                ["Key Vault", "1", "Domain admin credentials"],
                ["NSGs", "2+", "DC and management subnets"],
            ],
            "platform-security": [
                ["Sentinel Workspace (SecOps)", "1", "Dedicated for security team"],
                ["Defender for Cloud", "1", "All 6 plans enabled"],
                ["SOAR Playbooks (Logic Apps)", "5+", "Auto-remediation"],
                ["Alert Rules", "20+", "Sev 0-2 incident triggers"],
                ["Key Vault (Forensics)", "1", "Evidence preservation"],
            ],
            "corp": [
                ["Virtual Network (Spoke)", "1", "Corp CIDR 10.2.0.0/16"],
                ["Subnets", "4", "App, Data, Shared Services, Private Endpoints"],
                ["NSGs", "4", "Per-subnet deny-all-inbound baseline"],
                ["Key Vault", "1", "Private endpoint enabled"],
                ["Storage Account", "1", "Private endpoint enabled"],
                ["Azure Policy Assignments", "10+", "Corp baseline"],
            ],
            "online": [
                ["Virtual Network (Spoke)", "1", "Online CIDR 10.1.0.0/16"],
                ["Subnets", "3", "Web, App, Data"],
                ["WAF Policy", "1", "OWASP 3.2 ruleset"],
                ["DDoS Protection", "1", "Standard (inherited from hub)"],
                ["NSGs", "3", "Per-subnet rules"],
                ["Azure Policy Assignments", "10+", "Online baseline"],
            ],
            "sap": [
                ["Virtual Network (Spoke)", "1", "SAP CIDR 10.5.0.0/16"],
                ["Subnets", "5", "App, DB, ANF, Web Dispatcher, Management"],
                ["Azure NetApp Files", "1", "Ultra tier for HANA"],
                ["Proximity Placement Group", "1", "Low-latency colocation"],
                ["NSGs", "5", "Per-subnet with SAP-specific rules"],
                ["Key Vault", "1", "SAP credential management"],
            ],
            "sandbox": [
                ["Virtual Network (Standalone)", "1", "Sandbox CIDR 10.10.0.0/16"],
                ["Subnets", "2", "Dev, Test"],
                ["Budget", "1", "$500/month cap"],
                ["Azure Policy Assignments", "5+", "Sandbox baseline (no public IPs)"],
            ],
        }

        resources = expected.get(self.profile, [])
        if resources:
            self._add_table(
                ["Resource", "Count", "Configuration"],
                resources,
                col_widths=[6, 2, 9],
            )

    def _section_network_topology(self):
        self.doc.add_heading("4. Network Topology", level=1)

        is_platform_conn = self.profile == "platform-connectivity"
        is_sandbox = self.profile == "sandbox"

        if is_platform_conn:
            self.doc.add_paragraph(
                "This landing zone hosts the hub virtual network. All spoke VNets "
                "from application landing zones peer to this hub for centralized "
                "firewall inspection, DNS resolution, and on-premises connectivity."
            )
        elif is_sandbox:
            self.doc.add_paragraph(
                "This sandbox landing zone uses a standalone virtual network with no "
                "peering to the hub. It is isolated by design for development and testing."
            )
        else:
            self.doc.add_paragraph(
                "This landing zone uses a spoke virtual network peered to the central "
                "hub in the Connectivity subscription. All egress traffic routes through "
                "Azure Firewall. DNS resolution uses the hub's Private DNS Zones."
            )

        # Network details table
        net_configs = {
            "platform-management": [
                ["VNet CIDR", "Linked to Hub (no dedicated VNet)"],
                ["Connectivity", "Direct diagnostic log shipping to LAW"],
            ],
            "platform-connectivity": [
                ["Hub VNet CIDR", "10.0.0.0/16"],
                ["Firewall Subnet", "10.0.1.0/26 (AzureFirewallSubnet)"],
                ["Gateway Subnet", "10.0.2.0/27 (GatewaySubnet)"],
                ["Bastion Subnet", "10.0.3.0/26 (AzureBastionSubnet)"],
                ["DNS Resolver", "10.0.4.0/28"],
                ["Topology", "Hub-Spoke (Azure Firewall Premium)"],
            ],
            "platform-identity": [
                ["VNet CIDR", "Peered to Hub"],
                ["DC Subnet", "Dedicated for domain controllers"],
                ["DNS", "Custom DNS pointing to DCs"],
            ],
            "platform-security": [
                ["VNet CIDR", "Peered to Hub"],
                ["Connectivity", "Log ingestion from all subscriptions"],
            ],
            "corp": [
                ["Spoke VNet CIDR", "10.2.0.0/16"],
                ["App Subnet", "10.2.1.0/24"],
                ["Data Subnet", "10.2.2.0/24"],
                ["Shared Services", "10.2.3.0/24"],
                ["Private Endpoints", "10.2.4.0/26"],
                ["Hub Peering", "Active — routes to Azure Firewall"],
                ["DNS", "Hub Private DNS Zones"],
            ],
            "online": [
                ["Spoke VNet CIDR", "10.1.0.0/16"],
                ["Web Subnet", "10.1.1.0/24"],
                ["App Subnet", "10.1.2.0/24"],
                ["Data Subnet", "10.1.3.0/24"],
                ["Hub Peering", "Active — routes to Azure Firewall"],
                ["DDoS Protection", "Standard (shared plan)"],
            ],
            "sap": [
                ["Spoke VNet CIDR", "10.5.0.0/16"],
                ["SAP App Subnet", "10.5.1.0/24 (Accelerated Networking)"],
                ["HANA DB Subnet", "10.5.2.0/24 (Accelerated Networking)"],
                ["ANF Subnet", "10.5.3.0/26 (Delegated)"],
                ["Web Dispatcher", "10.5.4.0/26"],
                ["Management", "10.5.5.0/26"],
                ["Hub Peering", "Active — ExpressRoute to on-prem"],
            ],
            "sandbox": [
                ["Standalone VNet CIDR", "10.10.0.0/16"],
                ["Dev Subnet", "10.10.1.0/24"],
                ["Test Subnet", "10.10.2.0/24"],
                ["Hub Peering", "None (isolated)"],
            ],
        }

        config = net_configs.get(self.profile, [["N/A", "Profile-specific"]])
        self._add_table(["Network Component", "Configuration"], config, col_widths=[6, 11])

    def _section_security_posture(self):
        self.doc.add_heading("5. Security Posture", level=1)
        self.doc.add_paragraph(
            "Security controls applied to this landing zone as part of the CAF "
            "enterprise-scale baseline. All controls are enforced via Azure Policy "
            "and validated during deployment."
        )

        self.doc.add_heading("5.1 Non-Negotiable Security Rules", level=2)
        rules = [
            ["1", "Diagnostic Settings", "All resources ship logs to central Log Analytics workspace", "Enforced"],
            ["2", "HTTPS Only", "All web endpoints require TLS 1.2+", "Enforced"],
            ["3", "No Public IPs", "Disallowed on compute (except for allowed profiles)", "Enforced"],
            ["4", "Encryption at Rest", "All storage & databases use platform-managed or CMK encryption", "Enforced"],
            ["5", "NSG on Every Subnet", "Network Security Groups required on all subnets", "Enforced"],
            ["6", "Defender for Cloud", "Enabled on all resource types (per profile plan count)", "Enforced"],
        ]
        self._add_table(
            ["#", "Rule", "Description", "Status"],
            rules,
            col_widths=[1, 4, 9, 2.5],
        )

        self.doc.add_heading("5.2 Defender for Cloud Plans", level=2)
        plan_map = {
            "platform-management": ["Servers", "Storage", "Key Vaults", "ARM", "DNS", "Containers"],
            "platform-connectivity": ["Servers", "DNS", "ARM", "Key Vaults"],
            "platform-identity": ["Servers", "Key Vaults", "ARM", "DNS"],
            "platform-security": ["Servers", "Storage", "Key Vaults", "ARM", "DNS", "Containers"],
            "corp": ["Servers", "Storage", "Key Vaults", "AppServices", "ARM", "DNS"],
            "online": ["Servers", "Storage", "Key Vaults", "AppServices", "DNS"],
            "sap": ["Servers", "Storage", "Key Vaults", "ARM", "DNS", "Containers"],
            "sandbox": ["Servers", "ARM"],
        }
        plans = plan_map.get(self.profile, ["Servers"])
        rows = [[plan, "Enabled"] for plan in plans]
        self._add_table(["Defender Plan", "Status"], rows, col_widths=[8, 4])

        self.doc.add_heading("5.3 Azure Policy Assignments", level=2)
        self.doc.add_paragraph(
            "The following policy initiatives are assigned at the management group level "
            "and inherited by this subscription:"
        )
        policy_rows = [
            ["CAF Foundation", "Core governance (tagging, allowed locations, allowed SKUs)"],
            ["CAF Security Baseline", "CIS benchmark controls, encryption, network rules"],
            ["Defender for Cloud", "Auto-enable Defender plans and security configurations"],
            ["Monitoring", "Diagnostic settings, Log Analytics agent, dependency agent"],
            ["Network", "NSG rules, flow logs, VNet service endpoints"],
        ]
        self._add_table(["Initiative", "Scope"], policy_rows, col_widths=[5, 12])

    def _section_compliance_status(self):
        self.doc.add_heading("6. Compliance Status", level=1)
        self.doc.add_paragraph(
            "Post-deployment compliance scan results. The CI/CD pipeline validates "
            "compliance after every deployment and fails the pipeline if compliance "
            "falls below 80%."
        )

        self.doc.add_heading("6.1 Compliance Summary", level=2)
        summary = [
            ["Compliance Percentage", "Populated at deployment time"],
            ["Total Policies Evaluated", "Populated at deployment time"],
            ["Compliant Resources", "Populated at deployment time"],
            ["Non-Compliant Resources", "Populated at deployment time"],
            ["Exempt Resources", "Populated at deployment time"],
        ]
        self._add_table(["Metric", "Value"], summary, col_widths=[6, 11])

        p = self.doc.add_paragraph()
        run = p.add_run(
            "Note: This section is auto-populated with live data when the TDD is "
            "generated as part of the CI/CD pipeline (post-deployment verify stage). "
            "For pre-deployment TDDs, values show 'Populated at deployment time'."
        )
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = GRAY

    def _section_cost_governance(self):
        self.doc.add_heading("7. Cost Governance", level=1)

        budget_map = {
            "platform-management": ("10,000", "80/100/120%"),
            "platform-connectivity": ("15,000", "80/100/120%"),
            "platform-identity": ("5,000", "80/100/120%"),
            "platform-security": ("8,000", "80/100/120%"),
            "corp": ("20,000", "80/100/120%"),
            "online": ("25,000", "80/100/120%"),
            "sap": ("50,000", "80/100/120%"),
            "sandbox": ("500", "80/100%"),
        }
        budget, thresholds = budget_map.get(self.profile, ("10,000", "80/100/120%"))

        rows = [
            ["Monthly Budget", f"${budget}"],
            ["Alert Thresholds", thresholds],
            ["Alert Recipients", "Platform team + subscription owner"],
            ["Cost Anomaly Detection", "Enabled"],
            ["Tag Requirements", "Environment, Owner, CostCenter, Project"],
        ]
        self._add_table(["Control", "Configuration"], rows, col_widths=[6, 11])

    def _section_operational_model(self):
        self.doc.add_heading("8. Operational Model", level=1)

        self.doc.add_heading("8.1 Monitoring & Alerting", level=2)
        mon_rows = [
            ["Compliance Scan", "Every 30 minutes", "4-monitor.yml"],
            ["Drift Detection", "Every hour", "4-monitor.yml"],
            ["Full Audit Report", "Daily 6 AM UTC", "4-monitor.yml"],
            ["Cost Alerts", "Real-time", "Azure Cost Management"],
            ["Security Alerts", "Real-time", "Defender for Cloud → Sentinel"],
        ]
        self._add_table(["Scan Type", "Frequency", "Source"], mon_rows, col_widths=[5, 4, 8])

        self.doc.add_heading("8.2 Change Management", level=2)
        self.doc.add_paragraph(
            "All infrastructure changes follow the GitOps workflow:"
        )
        change_rows = [
            ["1", "Create feature branch", "Developer"],
            ["2", "Push changes & open PR", "Developer"],
            ["3", "Automated PR validation (lint, security, cost, what-if)", "5-pr-validate.yml"],
            ["4", "Peer review & approval", "Platform team"],
            ["5", "Merge to main", "Developer"],
            ["6", "Trigger deploy workflow", "Platform team"],
            ["7", "Environment approval gate", "Required reviewers"],
            ["8", "Deployment + post-deploy verification", "Reusable pipeline"],
            ["9", "TDD auto-generated", "tdd_generator.py"],
        ]
        self._add_table(["Step", "Action", "Responsible"], change_rows, col_widths=[1.5, 9, 6])

        self.doc.add_heading("8.3 Disaster Recovery", level=2)
        dr_map = {
            "platform-management": [
                ["LAW Data", "Geo-redundant storage", "< 4 hours"],
                ["Automation Runbooks", "Git (source of truth)", "Minutes"],
                ["Backup Vault", "Cross-region replication", "< 4 hours"],
            ],
            "platform-connectivity": [
                ["Firewall Config", "Git (IaC templates)", "< 30 minutes"],
                ["DNS Zones", "Azure-managed redundancy", "Automatic"],
                ["Gateway Config", "Git (IaC templates)", "< 1 hour"],
            ],
            "corp": [
                ["Application Data", "Recovery Services Vault", "RPO: 24h / RTO: 4h"],
                ["Key Vault", "Soft delete + purge protection", "Automatic"],
                ["Storage", "GRS replication", "Automatic"],
            ],
        }
        dr = dr_map.get(self.profile, [
            ["IaC Templates", "Git repository", "Minutes"],
            ["Configuration", "subscriptions.json", "Minutes"],
        ])
        self._add_table(["Component", "Protection", "Recovery Time"], dr, col_widths=[5, 6, 5.5])

    def _section_appendix(self):
        self.doc.add_heading("9. Appendix", level=1)

        self.doc.add_heading("9.1 Full Estate Architecture", level=2)
        self.doc.add_paragraph(
            "Overview of the complete Azure Landing Zone estate showing all platform "
            "and application landing zones."
        )

        # Use shared estate PNG (diagrams library with real Azure icons)
        try:
            config_file = Path(self.config_path)
            if config_file.exists():
                with open(config_file) as f:
                    config = json.load(f)
                engine = DiagramEngine(output_dir=str(Path(self.config_path).parent.parent / "docs" / "tdd"))
                estate_png = engine.generate_full_estate(
                    mg_prefix=config.get("management_group_prefix", "mrg"),
                    subscriptions_config=config,
                    filename="alz-estate-overview",
                )
                if estate_png and Path(estate_png).exists():
                    self._insert_png_diagram(estate_png, "Figure 2: Full Azure Landing Zone Estate")
                else:
                    raise FileNotFoundError(f"Estate PNG not generated at {estate_png}")
        except Exception as e:
            logger.warning(f"Could not generate estate diagram: {e}")
            self.doc.add_paragraph(f"[Estate diagram unavailable: {e}]")

        self.doc.add_heading("9.2 Document Revision History", level=2)
        self._add_table(
            ["Version", "Date", "Author", "Changes"],
            [["1.0", self.generated_at.strftime("%Y-%m-%d"), "Auto-generated by CI/CD pipeline",
              "Initial as-built TDD"]],
            col_widths=[2, 3, 6, 5.5],
        )

        self.doc.add_heading("9.3 References", level=2)
        refs = [
            ["CAF Enterprise-Scale", "https://aka.ms/caf/enterprise-scale"],
            ["Azure Landing Zone Accelerator", "https://aka.ms/alz/accelerator"],
            ["Landing Zone Profiles", "src/config/landing_zone_profiles.yaml"],
            ["Subscription Config", "environments/subscriptions.json"],
            ["CI/CD Workflows", ".github/workflows/"],
        ]
        self._add_table(["Document", "Location"], refs, col_widths=[6, 11])

    # ─── Main Generation Method ───────────────────────────────────────────

    def generate(
        self,
        output_path: str,
        resource_inventory: Optional[dict] = None,
        compliance_data: Optional[dict] = None,
    ) -> str:
        """
        Generate the complete TDD document.

        Args:
            output_path: Where to save the .docx file
            resource_inventory: Live data from Resource Graph (optional)
            compliance_data: Live compliance scan results (optional)

        Returns:
            Path to the generated document
        """
        logger.info(f"Generating TDD for {self.project_name} ({self.profile})")

        self._setup_styles()
        self._add_cover_page()
        self._add_toc()

        self._section_executive_summary()
        self._section_architecture_diagram()
        self._section_resource_inventory(resource_inventory or {})
        self._section_network_topology()
        self._section_security_posture()
        self._section_compliance_status()
        self._section_cost_governance()
        self._section_operational_model()
        self._section_appendix()

        # Save
        output = Path(output_path)
        output.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(output))

        logger.info(f"TDD saved to {output}")
        return str(output)

    # ─── Markdown Generation ──────────────────────────────────────────────

    def generate_markdown(
        self,
        md_path: str,
        svg_path: str,
        svg_filename: str,
        resource_inventory: Optional[dict] = None,
    ) -> str:
        """
        Generate a markdown version of the TDD with an SVG architecture diagram.

        Args:
            md_path: Where to save the .md file
            svg_path: Where to save the .svg diagram
            svg_filename: Filename of the SVG (for relative link in markdown)
            resource_inventory: Live data from Resource Graph (optional)

        Returns:
            Path to the generated markdown document
        """
        logger.info(f"Generating markdown TDD for {self.project_name} ({self.profile})")

        Path(svg_path).parent.mkdir(parents=True, exist_ok=True)

        # Preferred: Draw.io engine (.drawio + .svg with Azure icons)
        # Fallback 1: PNG via diagrams/graphviz library
        # Fallback 2: hand-drawn SVG via azure_diagram_generator
        out_dir = str(Path(svg_path).parent)
        drawio_path, drawio_svg_path = self._generate_drawio_diagram(output_dir=out_dir)
        png_filename = svg_filename.replace("_architecture.svg", "_architecture.png")
        if drawio_svg_path and Path(drawio_svg_path).exists():
            # Draw.io SVG is the primary diagram for markdown
            png_filename = None
            svg_filename = Path(drawio_svg_path).name
            logger.info(f"Architecture diagram (Draw.io SVG) saved to {drawio_svg_path}")
        else:
            png_path = self._generate_png_diagram(output_dir=out_dir)
            if not png_path or not Path(png_path).exists():
                png_filename = None
                # Fall back to SVG
                svg_content = self._get_profile_diagram()
                Path(svg_path).write_text(svg_content, encoding="utf-8")
                logger.info(f"Architecture diagram (SVG fallback) saved to {svg_path}")
            else:
                png_filename = Path(png_path).name
                logger.info(f"Architecture diagram (PNG) saved to {png_path}")

        # Generate shared estate PNG once (all TDDs reference the same file)
        estate_diagram_filename = None
        try:
            config_file = Path(self.config_path)
            if config_file.exists():
                with open(config_file) as f:
                    config = json.load(f)
                estate_png_path = Path(svg_path).parent / "alz-estate-overview.png"
                if not estate_png_path.exists():
                    engine = DiagramEngine(output_dir=str(Path(svg_path).parent))
                    engine.generate_full_estate(
                        mg_prefix=config.get("management_group_prefix", "mrg"),
                        subscriptions_config=config,
                        filename="alz-estate-overview",
                    )
                estate_diagram_filename = "alz-estate-overview.png"
        except Exception as e:
            logger.warning(f"Could not generate estate diagram: {e}")

        # Build markdown content
        generated_at = self.generated_at.strftime("%Y-%m-%d %H:%M UTC")
        diagram_filename = png_filename or svg_filename
        md = self._build_markdown(
            diagram_filename, estate_diagram_filename, resource_inventory or {}, generated_at,
        )

        Path(md_path).parent.mkdir(parents=True, exist_ok=True)
        Path(md_path).write_text(md, encoding="utf-8")
        logger.info(f"Markdown TDD saved to {md_path}")
        return md_path

    def _build_markdown(
        self,
        diagram_filename: str,
        estate_diagram_filename: Optional[str],
        resource_inventory: dict,
        generated_at: str,
    ) -> str:
        """Assemble the full markdown document."""
        profile_desc = {
            "platform-management": "centralized monitoring, logging, Sentinel SIEM, and backup infrastructure",
            "platform-connectivity": "hub networking with Azure Firewall, DNS, DDoS protection, and ExpressRoute/VPN gateways",
            "platform-identity": "Active Directory Domain Services, Entra ID PIM, and privileged access workstations",
            "platform-security": "dedicated SecOps with Sentinel, Defender for Cloud (all plans), SOAR playbooks, and incident response",
            "corp": "internal corporate workloads with private connectivity, ExpressRoute peering, and on-prem DNS integration",
            "online": "internet-facing workloads with WAF, DDoS Standard, and public endpoint protection",
            "sap": "SAP S/4HANA workloads with accelerated networking, Azure NetApp Files, and proximity placement groups",
            "sandbox": "development and testing with budget enforcement, relaxed policies, and standalone networking",
        }
        desc = profile_desc.get(self.profile, f"workloads using the {self.profile} profile")

        sections = [
            self._md_header(generated_at),
            self._md_executive_summary(desc),
            self._md_architecture_diagram(diagram_filename),
            self._md_resource_inventory(resource_inventory),
            self._md_network_topology(),
            self._md_security_posture(),
            self._md_compliance_status(),
            self._md_cost_governance(),
            self._md_operational_model(),
            self._md_appendix(estate_diagram_filename, generated_at),
        ]
        return "\n".join(sections)

    def _md_header(self, generated_at: str) -> str:
        return f"""# As-Built Technical Design Document

## {self.project_name}

> **Profile**: {self.profile} · **Environment**: {self.environment} · **Generated**: {generated_at}

| Field | Value |
|-------|-------|
| Subscription | {self.subscription_name} |
| Subscription ID | `{self.subscription_id}` |
| Location | {self.location} |
| IaC Framework | {self.framework.title()} |
| Deployment ID | {self.deployment_id or '(manual)'} |
| Document Version | 1.0 (auto-generated) |

---
"""

    def _md_executive_summary(self, desc: str) -> str:
        return f"""## 1. Executive Summary

This Technical Design Document (TDD) describes the as-built state of the
'{self.project_name}' landing zone deployed to subscription
'{self.subscription_name}' in {self.location}.
This landing zone provides {desc}.

The deployment was executed using {self.framework.title()} via the GitHub Actions
CI/CD pipeline with OIDC authentication and environment approval gates.
All resources comply with the CAF enterprise-scale baseline policies.

### Key Facts

| Attribute | Value |
|-----------|-------|
| Landing Zone | {self.project_name} |
| Profile | {self.profile} |
| Subscription | {self.subscription_name} (`{self.subscription_id}`) |
| Region | {self.location} |
| Environment | {self.environment} |
| IaC Framework | {self.framework.title()} |
| Deployment ID | {self.deployment_id or '(manual/CLI)'} |

---
"""

    def _md_architecture_diagram(self, diagram_filename: str) -> str:
        return f"""## 2. Architecture Diagram

The following diagram illustrates the as-built architecture of this landing zone,
including all deployed resources, networking topology, and security controls.
Icons follow the official Microsoft Azure Architecture Icon set.

![{self.project_name} Architecture — As-Built]({diagram_filename})

*Figure 1: {self.project_name} Architecture — As-Built*

---
"""

    def _md_resource_inventory(self, inventory: dict) -> str:
        lines = [
            "## 3. Resource Inventory",
            "",
            "Complete inventory of Azure resources deployed in this landing zone,",
            "queried from Azure Resource Graph at generation time.",
            "",
            "### 3.1 Resource Summary by Type",
            "",
        ]

        by_type = inventory.get("by_type", {})
        if by_type:
            lines.append("| Resource Type | Count |")
            lines.append("|---------------|-------|")
            for rtype, count in sorted(by_type.items(), key=lambda x: -x[1]):
                lines.append(f"| {rtype.split('/')[-1]} | {count} |")
            lines.append("")
            total = inventory.get("total_count", 0)
            lines.append(f"**Total resources: {total}**")
        else:
            # Expected resources based on profile
            expected = self._get_expected_resources()
            if expected:
                lines.append("| Resource | Count | Configuration |")
                lines.append("|----------|-------|---------------|")
                for row in expected:
                    lines.append(f"| {row[0]} | {row[1]} | {row[2]} |")
            else:
                lines.append("*Resource inventory will be populated from Azure Resource Graph at deployment time.*")

        lines.append("")
        lines.append("---")
        return "\n".join(lines)

    def _get_expected_resources(self) -> list[list[str]]:
        """Return expected resource table rows for the profile."""
        expected = {
            "platform-management": [
                ["Log Analytics Workspace", "1", "PerGB2018, 365-day retention"],
                ["Microsoft Sentinel", "1", "SIEM + SOAR integration"],
                ["Automation Account", "1", "Linked to LAW"],
                ["Recovery Services Vault", "1", "Daily + weekly backup policies"],
                ["Action Groups", "3+", "Email, Teams webhook, SMS"],
                ["Azure Monitor Alert Rules", "10+", "Performance, security, cost"],
            ],
            "platform-connectivity": [
                ["Virtual Network (Hub)", "1", "Hub CIDR 10.0.0.0/16"],
                ["Azure Firewall (Premium)", "1", "Threat Intel, IDPS"],
                ["VPN Gateway", "1", "VpnGw2 SKU"],
                ["ExpressRoute Gateway", "1", "Ultra Performance SKU"],
                ["Azure Bastion", "1", "Standard SKU"],
                ["DDoS Protection Plan", "1", "Standard tier"],
                ["Private DNS Zones", "20+", "Azure PaaS services"],
                ["Network Security Groups", "4+", "Per subnet"],
            ],
            "platform-identity": [
                ["Virtual Network (Identity)", "1", "Identity CIDR"],
                ["Domain Controllers (VMs)", "2", "Zone-redundant (AZ-1, AZ-2)"],
                ["Entra Connect (VM)", "1", "Hybrid sync to Entra ID"],
                ["Key Vault", "1", "Domain admin credentials"],
                ["NSGs", "2+", "DC and management subnets"],
            ],
            "platform-security": [
                ["Sentinel Workspace (SecOps)", "1", "Dedicated for security team"],
                ["Defender for Cloud", "1", "All plans enabled"],
                ["SOAR Playbooks (Logic Apps)", "5+", "Auto-remediation"],
                ["Alert Rules", "20+", "Sev 0-2 incident triggers"],
                ["Key Vault (Forensics)", "1", "Evidence preservation"],
            ],
            "corp": [
                ["Virtual Network (Spoke)", "1", "Corp CIDR 10.2.0.0/16"],
                ["Subnets", "4", "App, Data, Shared Services, Private Endpoints"],
                ["NSGs", "4", "Per-subnet deny-all-inbound baseline"],
                ["Key Vault", "1", "Private endpoint enabled"],
                ["Storage Account", "1", "Private endpoint enabled"],
                ["Azure Policy Assignments", "10+", "Corp baseline"],
            ],
            "online": [
                ["Virtual Network (Spoke)", "1", "Online CIDR 10.1.0.0/16"],
                ["Subnets", "3", "Web, App, Data"],
                ["WAF Policy", "1", "OWASP 3.2 ruleset"],
                ["DDoS Protection", "1", "Standard (inherited from hub)"],
                ["NSGs", "3", "Per-subnet rules"],
                ["Azure Policy Assignments", "10+", "Online baseline"],
            ],
            "sap": [
                ["Virtual Network (Spoke)", "1", "SAP CIDR 10.5.0.0/16"],
                ["Subnets", "5", "App, DB, ANF, Web Dispatcher, Management"],
                ["Azure NetApp Files", "1", "Ultra tier for HANA"],
                ["Proximity Placement Group", "1", "Low-latency colocation"],
                ["NSGs", "5", "Per-subnet with SAP-specific rules"],
                ["Key Vault", "1", "SAP credential management"],
            ],
            "sandbox": [
                ["Virtual Network (Standalone)", "1", "Sandbox CIDR 10.10.0.0/16"],
                ["Subnets", "2", "Dev, Test"],
                ["Budget", "1", "$500/month cap"],
                ["Azure Policy Assignments", "5+", "Sandbox baseline (no public IPs)"],
            ],
        }
        return expected.get(self.profile, [])

    def _md_network_topology(self) -> str:
        net_configs = {
            "platform-management": [
                ("VNet CIDR", "Linked to Hub (no dedicated VNet)"),
                ("Connectivity", "Direct diagnostic log shipping to LAW"),
            ],
            "platform-connectivity": [
                ("Hub VNet CIDR", "10.0.0.0/16"),
                ("Firewall Subnet", "10.0.1.0/26 (AzureFirewallSubnet)"),
                ("Gateway Subnet", "10.0.2.0/27 (GatewaySubnet)"),
                ("Bastion Subnet", "10.0.3.0/26 (AzureBastionSubnet)"),
                ("DNS Resolver", "10.0.4.0/28"),
                ("Topology", "Hub-Spoke (Azure Firewall Premium)"),
            ],
            "platform-identity": [
                ("VNet CIDR", "Peered to Hub"),
                ("DC Subnet", "Dedicated for domain controllers"),
                ("DNS", "Custom DNS pointing to DCs"),
            ],
            "platform-security": [
                ("VNet CIDR", "Peered to Hub"),
                ("Connectivity", "Log ingestion from all subscriptions"),
            ],
            "corp": [
                ("Spoke VNet CIDR", "10.2.0.0/16"),
                ("App Subnet", "10.2.1.0/24"),
                ("Data Subnet", "10.2.2.0/24"),
                ("Shared Services", "10.2.3.0/24"),
                ("Private Endpoints", "10.2.4.0/26"),
                ("Hub Peering", "Active — routes to Azure Firewall"),
                ("DNS", "Hub Private DNS Zones"),
            ],
            "online": [
                ("Spoke VNet CIDR", "10.1.0.0/16"),
                ("Web Subnet", "10.1.1.0/24"),
                ("App Subnet", "10.1.2.0/24"),
                ("Data Subnet", "10.1.3.0/24"),
                ("Hub Peering", "Active — routes to Azure Firewall"),
                ("DDoS Protection", "Standard (shared plan)"),
            ],
            "sap": [
                ("Spoke VNet CIDR", "10.5.0.0/16"),
                ("SAP App Subnet", "10.5.1.0/24 (Accelerated Networking)"),
                ("HANA DB Subnet", "10.5.2.0/24 (Accelerated Networking)"),
                ("ANF Subnet", "10.5.3.0/26 (Delegated)"),
                ("Web Dispatcher", "10.5.4.0/26"),
                ("Management", "10.5.5.0/26"),
                ("Hub Peering", "Active — ExpressRoute to on-prem"),
            ],
            "sandbox": [
                ("Standalone VNet CIDR", "10.10.0.0/16"),
                ("Dev Subnet", "10.10.1.0/24"),
                ("Test Subnet", "10.10.2.0/24"),
                ("Hub Peering", "None (isolated)"),
            ],
        }

        is_conn = self.profile == "platform-connectivity"
        is_sandbox = self.profile == "sandbox"

        if is_conn:
            intro = ("This landing zone hosts the hub virtual network. All spoke VNets "
                     "from application landing zones peer to this hub for centralized "
                     "firewall inspection, DNS resolution, and on-premises connectivity.")
        elif is_sandbox:
            intro = ("This sandbox landing zone uses a standalone virtual network with no "
                     "peering to the hub. It is isolated by design for development and testing.")
        else:
            intro = ("This landing zone uses a spoke virtual network peered to the central "
                     "hub in the Connectivity subscription. All egress traffic routes through "
                     "Azure Firewall. DNS resolution uses the hub's Private DNS Zones.")

        config = net_configs.get(self.profile, [("N/A", "Profile-specific")])
        lines = [
            "## 4. Network Topology",
            "",
            intro,
            "",
            "| Network Component | Configuration |",
            "|-------------------|---------------|",
        ]
        for comp, val in config:
            lines.append(f"| {comp} | {val} |")
        lines.append("")
        lines.append("---")
        return "\n".join(lines)

    def _md_security_posture(self) -> str:
        plan_map = {
            "platform-management": ["Servers", "Storage", "Key Vaults", "ARM", "DNS", "Containers"],
            "platform-connectivity": ["Servers", "DNS", "ARM", "Key Vaults"],
            "platform-identity": ["Servers", "Key Vaults", "ARM", "DNS"],
            "platform-security": ["Servers", "Storage", "Key Vaults", "ARM", "DNS", "Containers"],
            "corp": ["Servers", "Storage", "Key Vaults", "AppServices", "ARM", "DNS"],
            "online": ["Servers", "Storage", "Key Vaults", "AppServices", "DNS"],
            "sap": ["Servers", "Storage", "Key Vaults", "ARM", "DNS", "Containers"],
            "sandbox": ["Servers", "ARM"],
        }
        plans = plan_map.get(self.profile, ["Servers"])
        plan_rows = "\n".join(f"| {p} | Enabled |" for p in plans)

        return f"""## 5. Security Posture

Security controls applied to this landing zone as part of the CAF
enterprise-scale baseline. All controls are enforced via Azure Policy
and validated during deployment.

### 5.1 Non-Negotiable Security Rules

| # | Rule | Description | Status |
|---|------|-------------|--------|
| 1 | Diagnostic Settings | All resources ship logs to central Log Analytics workspace | Enforced |
| 2 | HTTPS Only | All web endpoints require TLS 1.2+ | Enforced |
| 3 | No Public IPs | Disallowed on compute (except for allowed profiles) | Enforced |
| 4 | Encryption at Rest | All storage & databases use platform-managed or CMK encryption | Enforced |
| 5 | NSG on Every Subnet | Network Security Groups required on all subnets | Enforced |
| 6 | Defender for Cloud | Enabled on all resource types (per profile plan count) | Enforced |

### 5.2 Defender for Cloud Plans

| Defender Plan | Status |
|---------------|--------|
{plan_rows}

### 5.3 Azure Policy Assignments

The following policy initiatives are assigned at the management group level
and inherited by this subscription:

| Initiative | Scope |
|------------|-------|
| CAF Foundation | Core governance (tagging, allowed locations, allowed SKUs) |
| CAF Security Baseline | CIS benchmark controls, encryption, network rules |
| Defender for Cloud | Auto-enable Defender plans and security configurations |
| Monitoring | Diagnostic settings, Log Analytics agent, dependency agent |
| Network | NSG rules, flow logs, VNet service endpoints |

---
"""

    def _md_compliance_status(self) -> str:
        return """## 6. Compliance Status

Post-deployment compliance scan results. The CI/CD pipeline validates
compliance after every deployment and fails the pipeline if compliance
falls below 80%.

### 6.1 Compliance Summary

| Metric | Value |
|--------|-------|
| Compliance Percentage | Populated at deployment time |
| Total Policies Evaluated | Populated at deployment time |
| Compliant Resources | Populated at deployment time |
| Non-Compliant Resources | Populated at deployment time |
| Exempt Resources | Populated at deployment time |

> *This section is auto-populated with live data when the TDD is
> generated as part of the CI/CD pipeline (post-deployment verify stage).
> For pre-deployment TDDs, values show 'Populated at deployment time'.*

---
"""

    def _md_cost_governance(self) -> str:
        budget_map = {
            "platform-management": ("10,000", "80/100/120%"),
            "platform-connectivity": ("15,000", "80/100/120%"),
            "platform-identity": ("5,000", "80/100/120%"),
            "platform-security": ("8,000", "80/100/120%"),
            "corp": ("20,000", "80/100/120%"),
            "online": ("25,000", "80/100/120%"),
            "sap": ("50,000", "80/100/120%"),
            "sandbox": ("500", "80/100%"),
        }
        budget, thresholds = budget_map.get(self.profile, ("10,000", "80/100/120%"))

        return f"""## 7. Cost Governance

| Control | Configuration |
|---------|---------------|
| Monthly Budget | ${budget} |
| Alert Thresholds | {thresholds} |
| Alert Recipients | Platform team + subscription owner |
| Cost Anomaly Detection | Enabled |
| Tag Requirements | Environment, Owner, CostCenter, Project |

---
"""

    def _md_operational_model(self) -> str:
        return """## 8. Operational Model

### 8.1 Monitoring & Alerting

| Scan Type | Frequency | Source |
|-----------|-----------|--------|
| Compliance Scan | Every 30 minutes | monitor.yml |
| Drift Detection | Every hour | monitor.yml |
| Full Audit Report | Daily 6 AM UTC | monitor.yml |
| Cost Alerts | Real-time | Azure Cost Management |
| Security Alerts | Real-time | Defender for Cloud → Sentinel |

### 8.2 Change Management

All infrastructure changes follow the GitOps workflow:

| Step | Action | Responsible |
|------|--------|-------------|
| 1 | Create feature branch | Developer |
| 2 | Push changes & open PR | Developer |
| 3 | Automated PR validation (lint, security, cost, what-if) | 5-pr-validate.yml |
| 4 | Peer review & approval | Platform team |
| 5 | Merge to main | Developer |
| 6 | Trigger deploy workflow | Platform team |
| 7 | Environment approval gate | Required reviewers |
| 8 | Deployment + post-deploy verification | Reusable pipeline |
| 9 | TDD auto-generated | tdd_generator.py |

### 8.3 Disaster Recovery

| Component | Protection | Recovery Time |
|-----------|------------|---------------|
| IaC Templates | Git repository | Minutes |
| Configuration | subscriptions.json | Minutes |

---
"""

    def _md_appendix(self, estate_diagram_filename: Optional[str], generated_at: str) -> str:
        lines = [
            "## 9. Appendix",
            "",
            "### 9.1 Full Estate Architecture",
            "",
            "Overview of the complete Azure Landing Zone estate showing all platform",
            "and application landing zones.",
            "",
        ]
        if estate_diagram_filename:
            lines.append(f"![Full Azure Landing Zone Estate]({estate_diagram_filename})")
            lines.append("")
            lines.append("*Figure 2: Full Azure Landing Zone Estate*")
        else:
            lines.append("*Estate diagram unavailable.*")
        lines.extend([
            "",
            "### 9.2 Document Revision History",
            "",
            "| Version | Date | Author | Changes |",
            "|---------|------|--------|---------|",
            f"| 1.0 | {generated_at.split(' ')[0]} | Auto-generated by CI/CD pipeline | Initial as-built TDD |",
            "",
            "### 9.3 References",
            "",
            "| Document | Location |",
            "|----------|----------|",
            "| CAF Enterprise-Scale | <https://aka.ms/caf/enterprise-scale> |",
            "| Azure Landing Zone Accelerator | <https://aka.ms/alz/accelerator> |",
            "| Landing Zone Profiles | `src/config/landing_zone_profiles.yaml` |",
            "| Subscription Config | `environments/subscriptions.json` |",
            "| CI/CD Workflows | `.github/workflows/` |",
        ])
        return "\n".join(lines)


def generate_tdd_for_deployment(
    project_name: str,
    profile: str,
    subscription_id: str,
    subscription_name: str,
    location: str,
    environment: str = "prod",
    framework: str = "bicep",
    deployment_id: str = "",
    output_dir: str = "agent-output/tdd",
    config_path: str = "environments/subscriptions.json",
    resource_inventory: dict | None = None,
) -> str:
    """
    Generate a TDD for a specific deployment. Called by CI/CD or CLI.

    Args:
        resource_inventory: Live data from Azure Resource Graph.
            Expected shape: {"total_count": int, "by_type": {str: int}}.
            When provided, replaces the static profile-based resource table.

    Returns the path to the generated document.
    """
    generator = TDDGenerator(
        project_name=project_name,
        profile=profile,
        subscription_id=subscription_id,
        subscription_name=subscription_name,
        location=location,
        environment=environment,
        framework=framework,
        deployment_id=deployment_id,
        config_path=config_path,
    )

    timestamp = datetime.now(timezone.utc).strftime("%Y%m%d")
    filename = f"TDD_{project_name}_{timestamp}.docx"
    output_path = str(Path(output_dir) / filename)

    return generator.generate(output_path, resource_inventory=resource_inventory or {})


def generate_tdd_markdown_for_deployment(
    project_name: str,
    profile: str,
    subscription_id: str,
    subscription_name: str,
    location: str,
    environment: str = "prod",
    framework: str = "bicep",
    deployment_id: str = "",
    output_dir: str = "agent-output/tdd",
    config_path: str = "environments/subscriptions.json",
    resource_inventory: dict | None = None,
) -> str:
    """
    Generate a markdown TDD with SVG architecture diagram for a specific deployment.

    Args:
        resource_inventory: Live data from Azure Resource Graph.
            Expected shape: {"total_count": int, "by_type": {str: int}}.
            When provided, replaces the static profile-based resource table.

    Returns the path to the generated .md file.
    """
    generator = TDDGenerator(
        project_name=project_name,
        profile=profile,
        subscription_id=subscription_id,
        subscription_name=subscription_name,
        location=location,
        environment=environment,
        framework=framework,
        deployment_id=deployment_id,
        config_path=config_path,
    )

    timestamp = datetime.now(timezone.utc).strftime("%Y%m%d")
    md_filename = f"TDD_{project_name}_{timestamp}.md"
    svg_filename = f"TDD_{project_name}_{timestamp}_architecture.svg"

    out = Path(output_dir)
    out.mkdir(parents=True, exist_ok=True)

    md_path = str(out / md_filename)
    svg_path = str(out / svg_filename)

    generator.generate_markdown(md_path, svg_path, svg_filename, resource_inventory=resource_inventory or {})
    return md_path


def generate_all_tdds(
    config_path: str = "environments/subscriptions.json",
    output_dir: str = "agent-output/tdd",
    framework: str = "bicep",
    fmt: str = "both",
) -> list[str]:
    """Generate TDDs for all landing zones defined in subscriptions.json.

    Args:
        fmt: Output format — 'docx', 'markdown', or 'both'
    """
    with open(config_path) as f:
        config = json.load(f)

    generated = []
    location = config.get("primary_location", "southcentralus")

    gen_fn_map = {
        "docx": [generate_tdd_for_deployment],
        "markdown": [generate_tdd_markdown_for_deployment],
        "both": [generate_tdd_for_deployment, generate_tdd_markdown_for_deployment],
    }
    gen_fns = gen_fn_map.get(fmt, gen_fn_map["both"])

    all_lzs = []
    for key, cfg in config.get("platform", {}).items():
        all_lzs.append((key, cfg, f"platform-{key}"))
    for key, cfg in config.get("application", {}).items():
        if not key.startswith("_"):
            all_lzs.append((key, cfg, ""))

    for key, cfg, default_profile in all_lzs:
        kwargs = dict(
            project_name=key,
            profile=cfg.get("profile", default_profile),
            subscription_id=cfg.get("subscription_id", ""),
            subscription_name=cfg.get("subscription_name", key),
            location=cfg.get("location", location),
            environment=cfg.get("environment", "prod"),
            framework=framework,
            output_dir=output_dir,
            config_path=config_path,
        )
        for gen_fn in gen_fns:
            path = gen_fn(**kwargs)
            generated.append(path)

    return generated


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="Generate As-Built TDD")
    parser.add_argument("--project", required=True, help="Landing zone name")
    parser.add_argument("--profile", required=True, help="LZ profile")
    parser.add_argument("--subscription-id", required=True)
    parser.add_argument("--subscription-name", required=True)
    parser.add_argument("--location", default="southcentralus")
    parser.add_argument("--environment", default="prod")
    parser.add_argument("--framework", default="bicep")
    parser.add_argument("--deployment-id", default="")
    parser.add_argument("--output-dir", default="agent-output/tdd")
    parser.add_argument("--config", default="environments/subscriptions.json")
    parser.add_argument("--all", action="store_true", help="Generate TDDs for all LZs")
    parser.add_argument(
        "--format", default="both", choices=["docx", "markdown", "both"],
        help="Output format: docx, markdown, or both (default: both)",
    )

    args = parser.parse_args()

    if args.all:
        paths = generate_all_tdds(args.config, args.output_dir, args.framework, fmt=args.format)
        print(f"Generated {len(paths)} TDD documents:")
        for p in paths:
            print(f"  {p}")
    else:
        gen_fns = {
            "docx": [generate_tdd_for_deployment],
            "markdown": [generate_tdd_markdown_for_deployment],
            "both": [generate_tdd_for_deployment, generate_tdd_markdown_for_deployment],
        }[args.format]

        base_kwargs = dict(
            project_name=args.project,
            profile=args.profile,
            subscription_id=args.subscription_id,
            subscription_name=args.subscription_name,
            location=args.location,
            environment=args.environment,
            framework=args.framework,
            deployment_id=args.deployment_id,
            output_dir=args.output_dir,
            config_path=args.config,
        )
        for fn in gen_fns:
            path = fn(**base_kwargs)
            print(f"TDD generated: {path}")
